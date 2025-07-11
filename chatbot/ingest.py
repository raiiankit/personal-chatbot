import os
import google.generativeai as genai
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain.document_loaders import TextLoader
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.embeddings import Embeddings
import pandas as pd
from docx import Document as DocxDocument

load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

DATA_DIR = "data"
DB_DIR = "chatbot/db"
EMBED_MODEL = "models/embedding-001"


# --- Helper to embed using Gemini
def get_embedding(text):
    return genai.embed_content(
        model=EMBED_MODEL,
        content=text,
        task_type="retrieval_document",
        title="PersonalData"
    )["embedding"]


# --- Gemini-compatible embedding wrapper
class GeminiEmbedding(Embeddings):
    def embed_documents(self, texts):
        return [get_embedding(t) for t in texts]

    def embed_query(self, text):
        return get_embedding(text)


# --- Load and extract content from different file types
def load_file(path, file_name):
    ext = file_name.lower().split(".")[-1]
    full_path = os.path.join(path, file_name)

    if ext in {"txt", "md"}:
        with open(full_path, "r") as f:
            return f.read()

    elif ext == "pdf":
        loader = PyPDFLoader(full_path)
        pages = loader.load()
        return "\n".join([p.page_content for p in pages])

    elif ext == "docx":
        doc = DocxDocument(full_path)
        return "\n".join([para.text for para in doc.paragraphs])

    elif ext == "csv":
        df = pd.read_csv(full_path)
        return df.to_string(index=False)

    return ""


# --- Ingest personal data into FAISS DB
def ingest_documents():
    print("🔍 Reading files...")
    embedding_model = GeminiEmbedding()

    docs = []

    for file_name in os.listdir(DATA_DIR):
        if file_name.lower().endswith((".txt", ".md", ".pdf", ".docx", ".csv")):
            content = load_file(DATA_DIR, file_name)
            if not content.strip():
                continue

            chunks = RecursiveCharacterTextSplitter(
                chunk_size=500, chunk_overlap=100
            ).split_text(content)

            for chunk in chunks:
                docs.append(Document(page_content=chunk, metadata={"source": file_name}))

    print(f"📚 Total chunks: {len(docs)}")

    texts = [d.page_content for d in docs]
    metadatas = [d.metadata for d in docs]

    print("🔐 Creating embeddings via Gemini...")
    if os.path.exists(os.path.join(DB_DIR, "index.faiss")):
        db = FAISS.load_local(DB_DIR, embeddings=embedding_model, allow_dangerous_deserialization=True)
        db.add_texts(texts, metadatas=metadatas)
    else:
        db = FAISS.from_texts(texts=texts, embedding=embedding_model, metadatas=metadatas)

    print("💾 Saving FAISS vector DB...")
    db.save_local(DB_DIR)
    print(f"✅ Saved at {DB_DIR}")


if __name__ == "__main__":
    ingest_documents()
